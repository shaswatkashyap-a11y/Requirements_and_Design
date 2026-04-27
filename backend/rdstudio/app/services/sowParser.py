import re
import logging
from pathlib import Path
from dataclasses import dataclass, field

import pymupdf4llm
import pymupdf
from docx import Document

from app.services.sectionClassifier import SectionClassifier

logger = logging.getLogger(__name__)


# ── Data classes returned by the parser ──────────────────────

@dataclass
class ParsedSection:
    title: str
    content: str
    level: int
    section_type: str = "unknown"
    confidence: float = 0.0
    page_number: int | None = None


@dataclass
class ParsedTable:
    headers: list[str]
    rows: list[list[str]]
    num_rows: int = 0
    parent_section: str | None = None


@dataclass
class ParsedSOW:
    filename: str
    raw_text: str
    markdown: str
    sections: list[ParsedSection] = field(default_factory=list)
    tables: list[ParsedTable] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


# ── Main parser class ────────────────────────────────────────

class SOWParser:
    """
    Two-stage SOW parser:
      1. pymupdf4llm / python-docx extracts structured content
      2. SectionClassifier categorizes each section
    """

    def __init__(self, classifier: SectionClassifier):
        self.classifier = classifier

    def parse(self, file_path: str) -> ParsedSOW:
        path = Path(file_path)
        ext = path.suffix.lower()
        logger.info(f"Parsing SOW: {path.name}")

        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    # ── PDF parsing ──────────────────────────────────────────

    def _parse_pdf(self, path: Path) -> ParsedSOW:
        # single markdown string
        markdown = pymupdf4llm.to_markdown(
            str(path),
            show_progress=False,
            page_chunks=False,
        )

        # per-page chunks for page number tracking
        page_chunks = pymupdf4llm.to_markdown(
            str(path),
            show_progress=False,
            page_chunks=True,
        )

        raw_text = ""
        for chunk in page_chunks:
            raw_text += chunk.get("text", "") + "\n"

        tables = self._extract_pdf_tables(str(path))
        raw_sections = self._markdown_to_sections(markdown, page_chunks)
        metadata = self._extract_metadata(raw_text)

        # classify each section
        classified = self._classify_sections(raw_sections)

        # link tables to their parent sections
        self._link_tables_to_sections(tables, classified)

        logger.info(
            f"Parsed {path.name}: {len(classified)} sections, "
            f"{len(tables)} tables"
        )

        markdown = re.sub(r"\*\*==>.*?<==\*\*", "", markdown)

        return ParsedSOW(
            filename=path.name,
            raw_text=raw_text.strip(),
            markdown=markdown,
            sections=classified,
            tables=tables,
            metadata=metadata,
        )

    def _extract_pdf_tables(self, file_path: str) -> list[ParsedTable]:
        tables = []
        doc = pymupdf.open(file_path)

        for page in doc:
            found = page.find_tables()
            for table in found.tables:
                try:
                    data = table.extract()
                    if not data or len(data) < 2:
                        continue
                    headers = [str(h).strip() for h in data[0]]
                    rows = [
                        [str(c).strip() for c in row]
                        for row in data[1:]
                    ]
                    tables.append(ParsedTable(
                        headers=headers,
                        rows=rows,
                        num_rows=len(rows),
                    ))
                except Exception:
                    continue

        doc.close()
        return tables

    # ── DOCX parsing ─────────────────────────────────────────

    def _parse_docx(self, path: Path) -> ParsedSOW:
        doc = Document(str(path))
        md_lines: list[str] = []
        raw_lines: list[str] = []
        sections: list[ParsedSection] = []
        current_heading: str | None = None
        current_level: int = 1
        content_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            raw_lines.append(text)
            style = (para.style.name or "").lower()

            is_heading = "heading" in style
            if not is_heading and self._is_bold_paragraph(para) and len(text.split()) < 12:
                is_heading = True

            if is_heading:
                level = self._get_heading_level(style)
                md_lines.append(f"{'#' * level} {text}")

                if current_heading is not None:
                    sections.append(ParsedSection(
                        title=current_heading,
                        content="\n".join(content_parts).strip(),
                        level=current_level,
                    ))
                    content_parts = []

                current_heading = self._clean_title(text)
                current_level = level
            else:
                md_lines.append(text)
                content_parts.append(text)

        # last section
        if current_heading is not None:
            sections.append(ParsedSection(
                title=current_heading,
                content="\n".join(content_parts).strip(),
                level=current_level,
            ))

        tables = self._extract_docx_tables(doc)
        raw_text = "\n".join(raw_lines)
        markdown = "\n\n".join(md_lines)
        metadata = self._extract_metadata(raw_text)

        # fallback if no headings found
        if not sections:
            sections = self._markdown_to_sections(markdown, [])

        classified = self._classify_sections(sections)
        self._link_tables_to_sections(tables, classified)

        return ParsedSOW(
            filename=path.name,
            raw_text=raw_text,
            markdown=markdown,
            sections=classified,
            tables=tables,
            metadata=metadata,
        )

    def _extract_docx_tables(self, doc: Document) -> list[ParsedTable]:
        tables = []
        for table in doc.tables:
            rows_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)
            if len(rows_data) >= 2:
                tables.append(ParsedTable(
                    headers=rows_data[0],
                    rows=rows_data[1:],
                    num_rows=len(rows_data) - 1,
                ))
        return tables

    # ── Markdown → Sections ──────────────────────────────────

    def _markdown_to_sections(
        self, markdown: str, page_chunks: list
    ) -> list[ParsedSection]:
        sections: list[ParsedSection] = []
        current_heading: str | None = None
        current_level: int = 1
        content_parts: list[str] = []

        page_map = self._build_page_map(page_chunks)
        heading_re = re.compile(r"^(#{1,6})\s+(.+)$")
        char_offset = 0

        for line in markdown.split("\n"):
            m = heading_re.match(line.strip())
            if m:
                if current_heading is not None:
                    sections.append(ParsedSection(
                        title=current_heading,
                        content="\n".join(content_parts).strip(),
                        level=current_level,
                        page_number=self._page_at_offset(page_map, char_offset),
                    ))
                    content_parts = []

                current_level = len(m.group(1))
                current_heading = self._clean_title(m.group(2))
            else:
                stripped = line.strip()
                if stripped:
                    content_parts.append(stripped)

            char_offset += len(line) + 1

        if current_heading is not None:
            sections.append(ParsedSection(
                title=current_heading,
                content="\n".join(content_parts).strip(),
                level=current_level,
                page_number=self._page_at_offset(page_map, char_offset),
            ))

        # last resort: split by numbered headings or ALL CAPS
        if not sections and markdown.strip():
            sections = self._split_by_patterns(markdown)

        return sections

    def _split_by_patterns(self, text: str) -> list[ParsedSection]:
        sections: list[ParsedSection] = []
        current_heading: str | None = None
        current_level = 1
        content_parts: list[str] = []

        numbered_re = re.compile(r"^\s*(\d+\.[\d.]*)\s+(.+)")
        allcaps_re = re.compile(r"^[A-Z][A-Z\s&/,:\-]{5,}$")

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            is_heading = False
            title = stripped
            level = 1

            m = numbered_re.match(stripped)
            if m:
                level = min(len(m.group(1).rstrip(".").split(".")), 3)
                title = self._clean_title(stripped) 
                is_heading = True
            elif allcaps_re.match(stripped) and len(stripped.split()) < 10:
                is_heading = True
                title = self._clean_title(stripped)

            if is_heading:
                if current_heading is not None:
                    sections.append(ParsedSection(
                        title=current_heading,
                        content="\n".join(content_parts).strip(),
                        level=current_level,
                    ))
                    content_parts = []
                current_heading = title
                current_level = level
            else:
                content_parts.append(stripped)

        if current_heading is not None:
            sections.append(ParsedSection(
                title=current_heading,
                content="\n".join(content_parts).strip(),
                level=current_level,
            ))

        return sections

    # ── Classification bridge ────────────────────────────────

    def _classify_sections(
        self, sections: list[ParsedSection]
    ) -> list[ParsedSection]:
        if not sections:
            return sections

        # convert to dicts for the classifier
        section_dicts = [
            {"title": s.title, "content": s.content}
            for s in sections
        ]
        classified = self.classifier.classify_batch(section_dicts)

        # write results back onto the dataclasses
        for section, result in zip(sections, classified):
            section.section_type = result["section_type"]
            section.confidence = result["confidence"]

        return sections

    # ── Link tables to sections ──────────────────────────────

    @staticmethod
    def _link_tables_to_sections(
        tables: list[ParsedTable],
        sections: list[ParsedSection],
    ):
        """
        Simple heuristic: assign each table to the section
        that appears just before it in document order.
        Works because tables usually sit inside the section
        they belong to.
        """
        if not sections or not tables:
            return

        # for now, assign tables evenly across sections
        # a more robust approach would use page numbers
        section_titles = [s.title for s in sections]
        for i, table in enumerate(tables):
            # assign to last section that could contain this table
            idx = min(i, len(section_titles) - 1)
            table.parent_section = section_titles[idx]

    # ── Helpers ───────────────────────────────────────────────

    @staticmethod
    def _build_page_map(page_chunks: list) -> list[tuple[int, int]]:
        page_map = []
        offset = 0
        for i, chunk in enumerate(page_chunks):
            text = chunk.get("text", "")
            page_num = chunk.get("metadata", {}).get("page", i + 1)
            page_map.append((offset, page_num))
            offset += len(text) + 1
        return page_map

    @staticmethod
    def _page_at_offset(
        page_map: list[tuple[int, int]], offset: int
    ) -> int | None:
        if not page_map:
            return None
        for i in range(len(page_map) - 1, -1, -1):
            if offset >= page_map[i][0]:
                return page_map[i][1]
        return page_map[0][1] if page_map else None

    @staticmethod
    def _is_bold_paragraph(para) -> bool:
        if not para.runs:
            return False
        return all(run.bold for run in para.runs if run.text.strip())

    @staticmethod
    def _get_heading_level(style_name: str) -> int:
        m = re.search(r"heading\s*(\d+)", style_name)
        return int(m.group(1)) if m else 1

    @staticmethod
    def _extract_metadata(raw_text: str) -> dict:
        meta = {}
        patterns = {
            "project_name": r"(?i)project\s*(?:name|title)\s*[:\-–]\s*(.+)",
            "client": r"(?i)(?:client|customer|prepared\s+for)\s*[:\-–]\s*(.+)",
            "vendor": r"(?i)(?:vendor|prepared\s+by|contractor)\s*[:\-–]\s*(.+)",
            "date": r"(?i)(?:date|effective\s*date)\s*[:\-–]\s*(.+)",
            "version": r"(?i)version\s*[:\-–]\s*(.+)",
            "contract_number": r"(?i)(?:contract|agreement)\s*(?:no|number|#)\s*[:\-–]\s*(.+)",
        }
        for key, pattern in patterns.items():
            m = re.search(pattern, raw_text)
            if m:
                meta[key] = m.group(1).strip()[:200]
        return meta
    
    @staticmethod
    def _clean_title(raw_title: str) -> str:
        """Clean markdown noise and formatting artifacts from section titles.
        
        Examples:
            '**STATEMENT OF WORK**'          → 'Statement of Work'
            '## **Key Activities**'           → 'Key Activities'
            '1.2 CUSTOMER ASSUMPTIONS:'       → '1.2 Customer Assumptions'
            '**Out of Scope:**'               → 'Out of Scope'
            'SECTION 3. COMPENSATION.'        → 'Section 3. Compensation'
            '**Escalation Matrix - Marvell**' → 'Escalation Matrix - Marvell'
        """
        title = raw_title.strip()

        # remove markdown heading markers (## etc)
        title = re.sub(r"^#{1,6}\s*", "", title)

        # remove markdown bold/italic markers
        title = title.replace("**", "").replace("__", "")
        title = title.replace("*", "").replace("_", "")

        # remove trailing colons and periods
        title = re.sub(r"[:.]+\s*$", "", title)

        # remove leading/trailing whitespace again
        title = title.strip()

        # normalize ALL CAPS to Title Case
        # but preserve mixed case like "UI/UX" or "OCM"
        if title.isupper() and len(title) > 5:
            # keep section numbers intact: "1.2 CUSTOMER ASSUMPTIONS" → "1.2 Customer Assumptions"
            parts = title.split(" ", 1)
            if re.match(r"^\d+[\d.]*$", parts[0]) and len(parts) > 1:
                title = f"{parts[0]} {parts[1].title()}"
            elif title.startswith("SECTION"):
                title = title.title()
            else:
                title = title.title()

        return title
















