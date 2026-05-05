import io
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from ..models.artifact import Artifact
from ..models.generationRun import GenerationRun
from ..models.module import Module
from .document_instruction_loader import load_instructions
from .methodology_config_loader import load_methodology_config

# ── Colour palette ────────────────────────────────────────────────────────────
COLOUR_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)   # indigo-800
COLOUR_ACCENT  = RGBColor(0x64, 0x74, 0x8B)   # slate-500

# ── Logo ──────────────────────────────────────────────────────────────────────
LOGO_PATH = Path(__file__).parent.parent / "assets" / "jadelogo.png"

# ── Artifact type display labels ──────────────────────────────────────────────
ARTIFACT_LABELS = {
    "functional_req":    "Functional Requirements",
    "nonfunctional_req": "Non-Functional Requirements",
    "architecture":      "Architecture",
    "task":              "Implementation Tasks",
    "test_case":         "Test Cases",
    "risk_entry":        "Risk Register",
}

# ── Extra section → artifact type mapping ─────────────────────────────────────
# Sections listed here are populated by aggregating that artifact type across
# all modules. Sections NOT listed show a clean placeholder — never raw guidance.
EXTRA_SECTION_ARTIFACT_SOURCE: dict[str, str] = {
    # Cloud architecture sections
    "aws_architecture_overview":      "architecture",
    "gcp_architecture_overview":      "architecture",
    "azure_architecture_overview":    "architecture",
    # Custom dev API / data model sections
    "api_specification":              "architecture",
    "data_model_overview":            "architecture",
    "angular_module_structure":       "architecture",
    "background_tasks_overview":      "architecture",
    # AI/ML pipeline sections
    "ml_pipeline_architecture":       "architecture",
    "agent_architecture":             "architecture",
    "data_pipeline_architecture":     "architecture",
    # Salesforce object model
    "salesforce_object_model":        "architecture",
    # NFR-adjacent sections
    "governor_limits_analysis":       "nonfunctional_req",
    "model_performance_requirements": "nonfunctional_req",
}


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────

def build_document(db: Session, run_id: int) -> bytes:
    run: GenerationRun | None = db.query(GenerationRun).filter(GenerationRun.id == run_id).first()
    if not run:
        raise ValueError(f"GenerationRun {run_id} not found")

    modules: list[Module] = (
        db.query(Module)
        .filter(Module.generation_run_id == run_id)
        .order_by(Module.module_order)
        .all()
    )

    # Group artifacts: {module_id: {artifact_type: [Artifact, ...]}}
    artifacts_by_module: dict[int, dict[str, list[Artifact]]] = {m.id: {} for m in modules}
    all_artifacts: list[Artifact] = (
        db.query(Artifact)
        .filter(Artifact.module_id.in_([m.id for m in modules]))
        .all()
    )
    for art in all_artifacts:
        artifacts_by_module[art.module_id].setdefault(art.artifact_type, []).append(art)

    instructions = load_instructions(run.service_line_codes or [], db)

    doc = Document()
    _configure_page(doc)
    _apply_base_styles(doc)

    _add_cover(doc, run, instructions, db)
    _add_executive_summary(doc, run, instructions)
    _add_module_overview_table(doc, modules)
    _add_extra_sections(doc, instructions, "after", "module_overview", modules, artifacts_by_module)
    _add_detailed_requirements(doc, modules, artifacts_by_module, run.methodology, instructions, db)
    _add_nfr_section(doc, modules, artifacts_by_module)
    _add_extra_sections(doc, instructions, "before", "appendix", modules, artifacts_by_module)
    _add_appendix(doc, instructions)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# Page & style helpers
# ─────────────────────────────────────────────────────────────────────────────

def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.54)
    section.top_margin  = section.bottom_margin = Cm(2.54)


def _apply_base_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _set_cell_shading(cell, hex_colour: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = COLOUR_PRIMARY
    run.font.name = "Calibri"


def _para(doc: Document, text: str, bold: bool = False, italic: bool = False,
          size: int = 10) -> None:
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def _page_break(doc: Document) -> None:
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# Standard section builders
# ─────────────────────────────────────────────────────────────────────────────

def _add_cover(doc: Document, run: GenerationRun, instructions: dict, db: Session) -> None:
    config = _get_methodology_config(run.methodology, db)
    doc_title = config["document_title"]

    _page_break(doc)
    doc.add_paragraph()

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Cm(5))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_get_project_name(run))
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = COLOUR_PRIMARY
    r.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(doc_title)
    r2.font.size = Pt(16)
    r2.font.color.rgb = COLOUR_ACCENT
    r2.font.name = "Calibri"

    doc.add_paragraph()

    meta = [
        ("Methodology",   config.get("display_name", run.methodology or "—")),
        ("Service Lines", ", ".join(instructions["display_names"]) or "—"),
        ("Generated",     date.today().strftime("%B %d, %Y")),
        ("Run ID",        str(run.id)),
        ("Status",        run.status or "—"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        r_lbl.font.name = "Calibri"
        r_lbl.font.size = Pt(11)
        r_val = p.add_run(value)
        r_val.font.name = "Calibri"
        r_val.font.size = Pt(11)

    _page_break(doc)


def _add_executive_summary(doc: Document, run: GenerationRun, instructions: dict) -> None:
    _heading(doc, "1. Executive Summary", level=1)
    system_name = instructions.get("terminology", {}).get("system", "the system")

    _para(doc, (
        f"This document presents the software requirements for {system_name} as understood "
        f"from the Statement of Work (SOW) provided by the client. It captures the functional "
        f"requirements, non-functional requirements, module structure, risks, and test cases "
        f"derived through AI-assisted analysis."
    ))
    doc.add_paragraph()

    base    = (run.sow.raw_text if run.sow else "").strip()
    summary = base[:1200] + ("…" if len(base) > 1200 else "")
    _para(doc, "SOW Summary:", bold=True)
    _para(doc, summary or "No SOW content provided.")
    doc.add_paragraph()


def _add_module_overview_table(doc: Document, modules: list[Module]) -> None:
    _heading(doc, "2. Module Overview", level=1)
    _para(doc, "The following functional modules were identified from the SOW:")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, label in enumerate(["#", "Module", "Description"]):
        hdr_cells[i].text = label
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Calibri"
        _set_cell_shading(hdr_cells[i], "1E40AF")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, mod in enumerate(modules, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = mod.name or ""
        row[2].text = mod.description or ""
        for cell in row:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9)
        if idx % 2 == 0:
            for cell in row:
                _set_cell_shading(cell, "F1F5F9")

    doc.add_paragraph()


def _add_extra_sections(
    doc: Document,
    instructions: dict,
    placement: str,
    relative_to: str,
    modules: list[Module],
    artifacts_by_module: dict[int, dict[str, list[Artifact]]],
) -> None:
    for section in instructions.get("extra_sections", []):
        if section.get("placement") != placement or section.get("relative_to") != relative_to:
            continue

        _heading(doc, section["title"], level=1)
        section_id = section.get("id", "")
        art_source = EXTRA_SECTION_ARTIFACT_SOURCE.get(section_id)

        if art_source == "architecture":
            _render_architecture_summary(doc, modules, artifacts_by_module)

        elif art_source == "nonfunctional_req":
            _render_nfr_summary(doc, modules, artifacts_by_module)

        else:
            # No direct artifact mapping — show a clean placeholder
            _para(
                doc,
                "This section is to be completed with project-specific details during "
                "project initiation. Refer to the detailed requirements and architecture "
                "sections for the AI-generated analysis.",
                italic=True,
            )

        doc.add_paragraph()


def _render_architecture_summary(
    doc: Document,
    modules: list[Module],
    artifacts_by_module: dict[int, dict[str, list[Artifact]]],
) -> None:
    """Cross-module architecture component summary table."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, label in enumerate(["Module", "Component", "Technology"]):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.name = "Calibri"
        _set_cell_shading(hdr[i], "1E40AF")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    row_idx = 0
    for mod in modules:
        arts = artifacts_by_module.get(mod.id, {}).get("architecture", [])
        for art in arts:
            # Each artifact record = one component (flat content_json)
            comp = art.content_json or {}
            row = table.add_row().cells
            row[0].text = mod.name or ""
            row[1].text = comp.get("component_name", "")
            row[2].text = comp.get("technology_suggestion", "")
            for cell in row:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.name = "Calibri"
                        r.font.size = Pt(9)
            if row_idx % 2 == 0:
                for cell in row:
                    _set_cell_shading(cell, "F1F5F9")
            row_idx += 1

    if row_idx == 0:
        table._element.getparent().remove(table._element)
        _para(doc, "No architecture components were generated for this run.", italic=True)


def _render_nfr_summary(
    doc: Document,
    modules: list[Module],
    artifacts_by_module: dict[int, dict[str, list[Artifact]]],
) -> None:
    """Aggregate NFRs across all modules for extra sections like governor limits."""
    found_any = False
    for mod in modules:
        arts = artifacts_by_module.get(mod.id, {}).get("nonfunctional_req", [])
        if not arts:
            continue
        found_any = True
        _para(doc, mod.name or "", bold=True)
        for art in arts:
            _render_nfr(doc, art.content_json or {})
    if not found_any:
        _para(doc, "No non-functional requirements were generated for this run.", italic=True)


def _add_detailed_requirements(
    doc: Document,
    modules: list[Module],
    artifacts_by_module: dict[int, dict[str, list[Artifact]]],
    methodology: str | None,
    instructions: dict,
    db: Session,
) -> None:
    _heading(doc, "3. Detailed Requirements by Module", level=1)
    config = _get_methodology_config(methodology, db)

    for mod in modules:
        _heading(doc, mod.name or f"Module {mod.id}", level=2)
        if mod.description:
            _para(doc, mod.description, italic=True)
        doc.add_paragraph()

        mod_artifacts = artifacts_by_module.get(mod.id, {})

        for art_type in config["artifact_order"]:
            arts = mod_artifacts.get(art_type, [])
            if not arts:
                continue
            label = config["functional_req_heading"] if art_type == "functional_req" else ARTIFACT_LABELS.get(art_type, art_type)
            _heading(doc, label, level=3)
            for art in arts:
                # Each DB record is one flat item — pass content_json directly
                _render_artifact(doc, art_type, art.content_json or {})
            doc.add_paragraph()

        doc.add_paragraph()


def _add_nfr_section(
    doc: Document,
    modules: list[Module],
    artifacts_by_module: dict[int, dict[str, list[Artifact]]],
) -> None:
    _heading(doc, "4. Non-Functional Requirements", level=1)
    found_any = False
    for mod in modules:
        arts = artifacts_by_module.get(mod.id, {}).get("nonfunctional_req", [])
        if not arts:
            continue
        found_any = True
        _heading(doc, mod.name or f"Module {mod.id}", level=2)
        for art in arts:
            _render_nfr(doc, art.content_json or {})
    if not found_any:
        _para(doc, "No non-functional requirements were generated for this run.")
    doc.add_paragraph()


def _add_appendix(doc: Document, instructions: dict) -> None:
    _heading(doc, "Appendix", level=1)

    standards = instructions.get("compliance_standards", [])
    if standards:
        _heading(doc, "A. Applicable Compliance Standards", level=2)
        for std in standards:
            _para(doc, std["name"], bold=True)
            _para(doc, f"  {std.get('description', '')}")
        doc.add_paragraph()

    roles = instructions.get("roles", [])
    if roles:
        _heading(doc, "B. Key Roles & Responsibilities", level=2)
        for role in roles:
            _para(doc, role["name"], bold=True)
            _para(doc, f"  {role.get('description', '')}")
        doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Artifact renderers — content is one flat dict (one DB record = one item)
# ─────────────────────────────────────────────────────────────────────────────

def _render_artifact(doc: Document, art_type: str, content: Any) -> None:
    if art_type == "functional_req":
        _render_functional_reqs(doc, content)
    elif art_type == "nonfunctional_req":
        _render_nfr(doc, content)
    elif art_type == "architecture":
        _render_architecture(doc, content)
    elif art_type == "task":
        _render_tasks(doc, content)
    elif art_type == "test_case":
        _render_test_cases(doc, content)
    elif art_type == "risk_entry":
        _render_risks(doc, content)
    else:
        _para(doc, str(content))


def _render_functional_reqs(doc: Document, content: Any) -> None:
    if not isinstance(content, dict):
        return
    _para(doc, f"{content.get('req_id', '')} — {content.get('title', '')}", bold=True)
    if content.get("user_story"):
        _para(doc, f"  User Story: {content['user_story']}")
    if content.get("description"):
        _para(doc, f"  {content['description']}")
    criteria = content.get("acceptance_criteria", [])
    if isinstance(criteria, list) and criteria:
        _para(doc, "  Acceptance Criteria:", bold=True)
        for c in criteria:
            doc.add_paragraph(f"    • {c}", style="List Bullet")
    if content.get("priority"):
        _para(doc, f"  Priority: {content['priority']}")
    if content.get("source_section"):
        _para(doc, f"  SOW Reference: {content['source_section']}")
    doc.add_paragraph()


def _render_nfr(doc: Document, content: Any) -> None:
    if not isinstance(content, dict):
        return
    _para(doc, f"{content.get('req_id', '')} [{content.get('category', '')}] — {content.get('title', '')}", bold=True)
    if content.get("description"):
        _para(doc, f"  {content['description']}")
    if content.get("measurable_criteria"):
        _para(doc, f"  Measurable Criteria: {content['measurable_criteria']}")
    if content.get("priority"):
        _para(doc, f"  Priority: {content['priority']}")
    doc.add_paragraph()


def _render_architecture(doc: Document, content: Any) -> None:
    if not isinstance(content, dict):
        return
    _para(doc, content.get("component_name", "Component"), bold=True)
    if content.get("description"):
        _para(doc, f"  {content['description']}")
    if content.get("technology_suggestion"):
        _para(doc, f"  Technology: {content['technology_suggestion']}")
    interfaces = content.get("interfaces", [])
    if isinstance(interfaces, list) and interfaces:
        _para(doc, f"  Interfaces: {', '.join(str(i) for i in interfaces)}")
    entities = content.get("data_entities", [])
    if isinstance(entities, list) and entities:
        _para(doc, f"  Data Entities: {', '.join(str(e) for e in entities)}")
    doc.add_paragraph()


def _render_tasks(doc: Document, content: Any) -> None:
    if not isinstance(content, dict):
        return
    task_type  = content.get("task_type", "")
    type_label = f" [{task_type}]" if task_type else ""
    _para(doc, f"{content.get('task_id', '')}{type_label} — {content.get('title', '')}", bold=True)
    if content.get("description"):
        _para(doc, f"  {content['description']}")
    if content.get("estimated_hours"):
        _para(doc, f"  Estimated Hours: {content['estimated_hours']}")
    if content.get("linked_requirement_id"):
        _para(doc, f"  Linked Requirement: {content['linked_requirement_id']}")
    criteria = content.get("acceptance_criteria", [])
    if isinstance(criteria, list) and criteria:
        _para(doc, "  Acceptance Criteria:", bold=True)
        for c in criteria:
            doc.add_paragraph(f"    • {c}", style="List Bullet")
    doc.add_paragraph()


def _render_test_cases(doc: Document, content: Any) -> None:
    if not isinstance(content, dict):
        return
    _para(doc, f"{content.get('test_id', '')} — {content.get('title', '')}", bold=True)
    _para(doc, f"  Type: {content.get('test_type', '')}  |  Linked: {content.get('linked_requirement_id', '')}")
    preconds = content.get("preconditions", [])
    if isinstance(preconds, list) and preconds:
        _para(doc, f"  Preconditions: {'; '.join(str(p) for p in preconds)}")
    steps = content.get("steps", [])
    if isinstance(steps, list) and steps:
        _para(doc, "  Steps:", bold=True)
        for step in steps:
            doc.add_paragraph(f"    {step}", style="List Number")
    if content.get("expected_result"):
        _para(doc, f"  Expected Result: {content['expected_result']}")
    doc.add_paragraph()


def _render_risks(doc: Document, content: Any) -> None:
    if not isinstance(content, dict):
        return
    likelihood = content.get("likelihood", "")
    impact     = content.get("impact", "")
    badge      = f"  [{likelihood.upper()} likelihood / {impact.upper()} impact]" if likelihood else ""
    # Risks have no title field — risk_id + badge as the header
    _para(doc, f"{content.get('risk_id', '')}{badge}", bold=True)
    if content.get("description"):
        _para(doc, f"  {content['description']}")
    if content.get("mitigation"):
        _para(doc, f"  Mitigation: {content['mitigation']}")
    if content.get("owner"):
        _para(doc, f"  Owner: {content['owner']}")
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Utility helpers
# ─────────────────────────────────────────────────────────────────────────────

def _get_methodology_config(methodology: str | None, db: Session) -> dict:
    if not methodology:
        raise ValueError("Methodology is required to generate the document but was not set on this run.")
    return load_methodology_config(methodology, db)


def _get_project_name(run: GenerationRun) -> str:
    if hasattr(run, "project") and run.project and hasattr(run.project, "name"):
        return run.project.name
    return f"Project Requirements — Run #{run.id}"
