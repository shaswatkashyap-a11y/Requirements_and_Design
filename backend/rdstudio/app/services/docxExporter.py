import re
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.designRun import DesignRun, DesignArtifact
from app.models.project import Project

HLD_SECTION_ORDER = [
    "folder_structure",
    "component_structure",
    "design_patterns",
    "technology",
    "error_handling",
    "api_design",
    "database_design",
    "security_architecture",
]

HLD_SECTION_LABELS = {
    "folder_structure":      "Folder Structure",
    "component_structure":   "Component Structure",
    "design_patterns":       "Design Patterns",
    "technology":            "Technology Stack",
    "error_handling":        "Error Handling Strategy",
    "api_design":            "API Design",
    "database_design":       "Database and Data Model",
    "security_architecture": "Security Architecture",
}


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D1D5DB")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_paragraph(p, hex_color: str):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


# ── Inline markdown parser (bold + code) ─────────────────────────────────────

_INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def _add_inline(para, text: str, base_size: int = 10):
    for part in _INLINE_PATTERN.split(text):
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        else:
            run = para.add_run(part)
            run.font.size = Pt(base_size)


# ── Table renderer ────────────────────────────────────────────────────────────

_SEP_RE = re.compile(r"^[\|\-\:\s]+$")


def _parse_md_row(line: str) -> list[str]:
    return [c.strip() for c in line.split("|")[1:-1]]


def _add_md_table(doc, table_lines: list[str]):
    data = [l for l in table_lines if not _SEP_RE.match(l)]
    if not data:
        return

    headers = _parse_md_row(data[0])
    body    = [_parse_md_row(l) for l in data[1:]]
    ncols   = len(headers)
    nrows   = 1 + len(body)

    table = doc.add_table(rows=nrows, cols=ncols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _set_table_borders(table)

    # Header
    for ci, text in enumerate(headers):
        cell = table.rows[0].cells[ci]
        _shade_cell(cell, "EEF2FF")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x31, 0x27, 0x8A)

    # Body
    for ri, row_data in enumerate(body):
        bg = "FFFFFF" if ri % 2 == 0 else "F9FAFB"
        row = table.rows[ri + 1]
        for ci in range(ncols):
            cell = row.cells[ci]
            _shade_cell(cell, bg)
            p = cell.paragraphs[0]
            cell_text = row_data[ci] if ci < len(row_data) else ""
            _add_inline(p, cell_text, base_size=9)

    doc.add_paragraph()


# ── Block markdown renderer ───────────────────────────────────────────────────

def _add_markdown(doc, markdown: str):
    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            _shade_paragraph(p, "1F2937")
            p.paragraph_format.left_indent  = Inches(0.1)
            p.paragraph_format.right_indent = Inches(0.1)
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x34, 0xD3, 0x99)
            continue

        # Table block
        if line.startswith("|"):
            tbl_lines: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            _add_md_table(doc, tbl_lines)
            continue

        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # Bullet
        elif re.match(r"^[-*•] ", line):
            text = re.sub(r"^[-*•] ", "", line)
            p = doc.add_paragraph(style="List Bullet")
            # clear auto-added run, then add inline-formatted text
            for run in p.runs:
                run.text = ""
            _add_inline(p, text)

        # Empty line → small spacer
        elif line.strip() == "":
            pass

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)

        i += 1


# ── Public builder ────────────────────────────────────────────────────────────

def build_hld_docx(run: DesignRun, artifacts: list, project: Project) -> BytesIO:
    doc = Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # ── Cover page ────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(72)
    tr = title_p.add_run("High-Level Design Document")
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = RGBColor(0x31, 0x27, 0x8A)

    proj_p = doc.add_paragraph()
    proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = proj_p.add_run(project.name)
    pr.font.size = Pt(16)
    pr.bold = True

    if project.client_name:
        c_p = doc.add_paragraph()
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c_p.add_run(f"Client: {project.client_name}")
        cr.font.size = Pt(12)
        cr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()

    # ── Sections ──────────────────────────────────────────────────────────────
    artifact_map = {a.section_type: a for a in artifacts}

    for section_type in HLD_SECTION_ORDER:
        artifact = artifact_map.get(section_type)
        if not artifact or not artifact.content_markdown:
            continue

        label = HLD_SECTION_LABELS.get(section_type, section_type.replace("_", " ").title())
        doc.add_heading(label, level=1)
        _add_markdown(doc, artifact.content_markdown)
        doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
